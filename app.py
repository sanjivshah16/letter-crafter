import streamlit as st
import hashlib
from datetime import date
import io
from openai import OpenAI
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
import pandas as pd
import fitz  

# --- Page config ---
st.set_page_config(page_title="Letter Crafter (Public)", layout="wide")
st.title("📄 Letter Crafter")

# --- Password protection ---
def verify_password(pw: str) -> bool:
    return hashlib.sha256(pw.encode()).hexdigest() == st.secrets.get("password_hash", "")

if "authenticated" not in st.session_state:
    st.session_state.authenticated = False

if not st.session_state.authenticated:
    pw = st.text_input("Enter password", type="password")
    if pw and verify_password(pw):
        st.session_state.authenticated = True
        st.rerun()
    elif pw:
        st.error("Incorrect password.")
    st.stop()

# --- OpenAI client ---
client = OpenAI(api_key=st.secrets["openai_api_key"])

# --- File extractors ---
def extract_text_from_pdf(file):
    try:
        doc = fitz.open(stream=file.read(), filetype="pdf")
        return "\n".join([page.get_text() for page in doc])
    except Exception as e:
        return f"(Failed to extract text from PDF: {e})"


def extract_text_from_docx(file):
    doc = Document(file)
    return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])

def extract_text_from_xlsx(file):
    dfs = pd.read_excel(file, sheet_name=None)
    output = []
    for name, df in dfs.items():
        output.append(f"Sheet: {name}\n")
        output.append(df.head(20).to_string(index=False))
        output.append("\n")
    return "\n".join(output)

def prepare_file_context(files):
    previews = []
    for f in files:
        filename = f.name
        try:
            if filename.endswith(".pdf"):
                text = extract_text_from_pdf(f)
            elif filename.endswith(".docx"):
                text = extract_text_from_docx(f)
            elif filename.endswith(".xlsx"):
                text = extract_text_from_xlsx(f)
            else:
                text = f.read().decode(errors="ignore")
            preview = text[:2000]
            previews.append(f"{filename}:\n{preview}\n")
        except Exception as e:
            previews.append(f"{filename}: (Could not extract text: {e})")
    return "\n".join(previews)

# --- Inputs ---
st.subheader("📁 Upload Materials")
uploaded_files = st.file_uploader("Upload CVs, drafts, personal statements, etc.", accept_multiple_files=True)

st.subheader("👥 Describe Your Relationship")
relationship_text = st.text_area("How do you know the applicant? (1–2 sentences)", height=120)

addressee = st.text_input("Addressee (e.g., Admissions Committee)", "")
salutation = st.text_input("Salutation (e.g., Dear Committee)", "")
if not salutation.strip():
    salutation = "To Whom It May Concern"

letter_date = date.today().strftime("%B %d, %Y")
filename = st.text_input("Output filename (no extension)", value="recommendation_letter")

font_name = st.selectbox("Font", ["Arial", "Times New Roman", "Calibri", "Aptos"], index=0)
font_size = st.selectbox("Font size", [9, 10, 10.5, 11, 11.5, 12], index=3)

# --- Template Upload ---
"Or [Download a Sample Template](https://www.dropbox.com/scl/fi/on6f93fpzzqy3zbug595y/LOR_template.docx?rlkey=lyyufxhkfgd0zb0ayvtxrihxu&dl=1)",
template_file = st.file_uploader("Upload a .docx template with placeholders", type=["docx"])

st.markdown(
    "Or [Download a Sample Template](https://www.dropbox.com/scl/fi/on6f93fpzzqy3zbug595y/LOR_template.docx?rlkey=lyyufxhkfgd0zb0ayvtxrihxu&dl=1)",
    unsafe_allow_html=True
)

# --- Generate letter with GPT-4o ---
def generate_letter(relationship_text, files):
    system_prompt = (
        "You are Letter Crafter, an expert letter writer. You will receive a description of the recommender's "
        "relationship with the applicant and readable text previews of attached files (e.g., CVs, drafts, etc). "
        "Use this information to write the body of a polished recommendation letter. "
        "Do NOT include the date, salutation, or closing. Return only the letter body."
    )

    file_context = prepare_file_context(files)
    user_prompt = (
        f"My relationship to the applicant:\n{relationship_text}\n\n"
        f"Attached files:\n{file_context}\n\n"
        f"Please write a professional recommendation letter body only."
    )

    try:
        response = client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt}
            ],
            temperature=0.7,
            max_tokens=1000
        )
        return response.choices[0].message.content.strip()
    except Exception as e:
        st.error(f"Error generating letter: {e}")
        return None

# --- Generate button ---
if st.button("✍️ Generate Letter"):
    if not uploaded_files or not relationship_text.strip():
        st.warning("Please upload at least one file and describe your relationship.")
        st.stop()

    if not template_file:
        st.warning("Please upload a Word template.")
        st.stop()

    letter_body = generate_letter(relationship_text, uploaded_files)
    if letter_body:
        st.session_state.letter_text = letter_body
        st.session_state.addressee = addressee
        st.session_state.salutation = salutation
        st.session_state.date = letter_date
        st.session_state.template_file = template_file
        st.success("Letter body generated.")

# --- Template insertion ---
def replace_placeholders(doc, replacements):
    date_idx = None
    for idx, p in enumerate(doc.paragraphs):
        for placeholder, replacement in replacements.items():
            if placeholder in p.text:
                p.clear()
                run = p.add_run(replacement)
                run.font.name = font_name
                run.font.size = Pt(font_size)
                run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
                if placeholder == "<<Date>>":
                    date_idx = idx

    if date_idx is not None:
        empties = []
        i = date_idx + 1
        while i < len(doc.paragraphs):
            para = doc.paragraphs[i]
            if para.text.strip():
                break
            empties.append(para)
            i += 1
        for p in empties[1:]:
            try:
                p._element.getparent().remove(p._element)
            except Exception:
                pass

# --- Format & download letter ---
if "letter_text" in st.session_state:
    try:
        doc = Document(st.session_state.template_file)
        replacements = {
            "<<Date>>": st.session_state.date,
            "<<Addressee>>": st.session_state.addressee,
            "<<Salutation>>": st.session_state.salutation,
            "<<Enter text here>>": st.session_state.letter_text
        }
        replace_placeholders(doc, replacements)

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        st.download_button(
            label="Download Letter (DOCX)",
            data=buffer,
            file_name=f"{filename}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
    except Exception as e:
        st.error(f"Error formatting letter: {e}")
